import streamlit as st
import pandas as pd
from Bio import SeqIO
from Bio.Seq import Seq
from Bio.SeqRecord import SeqRecord
from Bio.SeqFeature import SeqFeature, FeatureLocation
from Bio.Restriction import RestrictionBatch
import os
from io import BytesIO, StringIO
import zipfile
from datetime import datetime
from docx import Document

st.set_page_config(page_title="Crypto Designer", page_icon="🧪", layout="wide")
st.title("🧪 Crypto Allele Designer")

# ── 파일 체크 ──────────────────────────────────────────────────────────────────
if not os.path.exists("genome.fasta") or not os.path.exists("annotation.gff"):
    st.error("❌ 'genome.fasta' 및 'annotation.gff' 파일이 필요합니다. 작업 폴더에 두 파일이 있는지 확인해 주세요.")
    st.stop()


# ── 캐시된 파일 로딩 ────────────────────────────────────────────────────────────
@st.cache_resource
def load_genome():
    return SeqIO.to_dict(SeqIO.parse("genome.fasta", "fasta"))

@st.cache_resource
def load_gff():
    with open("annotation.gff", "r") as f:
        return f.readlines()


# ── Word 매뉴얼 생성 ────────────────────────────────────────────────────────────
def generate_manual_word():
    doc = Document()
    doc.add_heading('Crypto Allele Designer 사용 매뉴얼', 0)

    doc.add_heading('1. 주의사항 (Notice)', level=1)
    doc.add_paragraph('• 교차 점검 필수: 본 도구는 연구 편의를 돕는 자동화 스크립트입니다. 출력된 최종 Allele(.gb) 서열과 프라이머 결합 위치는 실험 진행 전, SnapGene 등 시퀀스 뷰어 프로그램을 통해 반드시 직접 교차 점검(Cross-check)하시기 바랍니다.')
    doc.add_paragraph('• 이용 범위: 본 도구는 학술적 연구 목적으로만 자유롭게 이용 가능하며, 상업적 이용은 엄격히 금지됩니다.')

    doc.add_heading('2. 기본 파일 설정', level=1)
    doc.add_paragraph('프로그램 실행 경로 내에 다음 두 파일이 반드시 존재해야 정상 작동합니다.')
    doc.add_paragraph('1) genome.fasta : Reference Genome 전체 서열\n2) annotation.gff : 유전자 위치(CDS 등)를 포함한 Annotation 파일')
    doc.add_paragraph('결과물은 ZIP 파일로 다운로드되며, [GB → DNA 변환] 탭을 이용하면 생성된 .gb 파일을 SnapGene에서 열람하기 편한 .dna 파일로 일괄 변환할 수 있습니다.')

    doc.add_heading('3. Excel 일괄(Batch) 양식 작성법', level=1)
    doc.add_paragraph('사이드바에서 다운로드한 엑셀 양식을 활용하여 다수의 유전자를 한 번에 디자인할 수 있습니다.')
    doc.add_paragraph('[Primers 시트]\n설계할 모든 프라이머의 Gene ID, Primer Name, Sequence를 입력합니다.')
    doc.add_paragraph('[Probes_WT / Probes_MUT 시트]\nWT 및 Mutant 서열에서 Southern blot 등에 사용할 Probe 구간의 양 끝 프라이머 이름을 기입합니다. (프라이머의 Overlap 서열은 제외되고 실제 주형과 결합하는 구간만 Probe로 표시됩니다.)')
    doc.add_paragraph('[Enzymes 시트]\n지도에 표기할 제한효소가 있다면 쉼표(,)로 구분하여 기입합니다. (예: BamHI, EcoRV)')
    doc.add_paragraph('[Jobs 시트]\n디자인의 주요 실행 옵션을 설정하는 시트입니다.\n- Output Mode: wt (WT만), mut (Mutant만), both (둘 다 출력)\n- Insert Mode: insert (단순 삽입), replace (특정 구간 치환). Replace 모드 선택 시, 치환되어 삭제되는 구간에 포함된 Exon 영역도 자동으로 계산되어 제거됩니다.\n- Primer A Overlap: WT 서열이 아닌 삽입 마커 서열 쪽에 결합하는 프라이머의 5\' Overlap 서열을 기입합니다.\n- 삽입 마커: 업로드한 마커 GB 파일(예: NAT.gb) 내의 주석(Feature) 정보는 Mutant 서열 생성 시 기존 위치와 이름이 보존되어 자동 이식됩니다.')

    output = BytesIO()
    doc.save(output)
    return output.getvalue()


# ── 엑셀 양식 생성 ────────────────────────────────────────────────────────────
def generate_template():
    output = BytesIO()
    df_p = pd.DataFrame({
        'Gene ID': ['CNAG_03701', 'CNAG_03701'],
        'Primer Name': ['L1', 'L2'],
        'Sequence': ['GCTTGTTGGCTTTCAGATG', 'CACTCGAATCCTGCATGCGTTGCCTTTTCTGTCGCC'],
    })
    df_pb_wt = pd.DataFrame({
        'Gene ID': ['CNAG_03701'],
        'Probe Start Primer': ['L1'],
        'Probe End Primer': ['L2'],
    })
    df_pb_mut = pd.DataFrame({
        'Gene ID': ['CNAG_03701'],
        'Probe Start Primer': ['NAT_F'],
        'Probe End Primer': ['NAT_R'],
    })
    df_e = pd.DataFrame({
        'Gene ID': ['CNAG_03701'],
        'Enzymes': ['ClaI'],
    })
    df_jobs = pd.DataFrame({
        'Gene ID':            ['CNAG_03701'],
        'Output Mode':        ['both'],        
        'Insert Mode':        ['replace'],      
        'Primer A':           ['L2'],
        'Primer A Overlap':   ['CACTCGAATCCTGCATGC'],  
        'Primer B':           ['R1'],            
        'Primer B Overlap':   [''],
        'Insert GB Filename': ['PCTR4_NAT.gb'],
        'Output Filename':    ['CNAG_03701_NAT'],  
    })
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df_p.to_excel(writer,      index=False, sheet_name='Primers')
        df_pb_wt.to_excel(writer,  index=False, sheet_name='Probes_WT')
        df_pb_mut.to_excel(writer, index=False, sheet_name='Probes_MUT')
        df_e.to_excel(writer,      index=False, sheet_name='Enzymes')
        df_jobs.to_excel(writer,   index=False, sheet_name='Jobs')
    return output.getvalue()


# ── 사이드바 ────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.header("⚙️ Settings")
    flank_size = st.number_input("Flanking Region (bp)", value=5000, step=500)

    st.divider()
    st.subheader("📐 Topology")
    topology = st.radio(
        "Linear / Circular", options=["linear", "circular"], index=0,
        help="모든 출력 파일에 일괄 적용됩니다."
    )

    st.divider()
    st.subheader("🗂️ ZIP 구조")
    zip_structure = st.radio(
        "파일 정리 방식",
        options=["flat", "by_gene"],
        format_func=lambda x: (
            "📄 폴더 없이 전체 압축" if x == "flat"
            else "📁 유전자별 폴더로 정리"
        ),
        index=0,
        help="by_gene: 동일한 유전자의 WT 및 MUT 파일을 하나의 폴더로 묶어 압축합니다."
    )

    st.divider()
    st.subheader("📥 다운로드 센터")
    st.download_button("📂 엑셀 양식 다운로드", generate_template(), "Allele_Template.xlsx")
    st.download_button("📘 사용 매뉴얼 (Word)", generate_manual_word(), "Crypto_Allele_Designer_Manual.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# ── 탭 구성 ─────────────────────────────────────────────────────────────────────
tab_main, tab_conv = st.tabs(["🧬 Allele Designer", "🔄 GB → DNA 변환"])


# ════════════════════════════════════════════════════════════════════════════════
# 공통 분석 함수
# ════════════════════════════════════════════════════════════════════════════════

def get_primer_coords(sub_seq, p_list):
    p_idx = {}
    for p in p_list:
        full_seq = str(p['seq']).strip().replace(" ", "").upper()
        if not full_seq: continue
        overlap = str(p.get('overlap', '')).strip().replace(" ", "").upper()

        core_seq = full_seq[len(overlap):] if (overlap and full_seq.startswith(overlap)) else full_seq
        if not core_seq: continue

        core_start, core_end, strand = -1, -1, 1
        target_fwd = Seq(core_seq)
        idx_fwd = sub_seq.find(target_fwd)
        if idx_fwd != -1:
            core_start = idx_fwd; core_end = idx_fwd + len(core_seq); strand = 1
        else:
            target_rev = target_fwd.reverse_complement()
            idx_rev = sub_seq.find(target_rev)
            if idx_rev != -1:
                core_start = idx_rev; core_end = idx_rev + len(core_seq); strand = -1

        if core_start == -1: continue 

        full_start, full_end = core_start, core_end 
        target_full_fwd = Seq(full_seq)
        idx_full_fwd = sub_seq.find(target_full_fwd)
        if idx_full_fwd != -1:
            full_start = idx_full_fwd; full_end = idx_full_fwd + len(full_seq)
        else:
            target_full_rev = target_full_fwd.reverse_complement()
            idx_full_rev = sub_seq.find(target_full_rev)
            if idx_full_rev != -1:
                full_start = idx_full_rev; full_end = idx_full_rev + len(full_seq)

        p_idx[p['name']] = {
            'full_start': full_start, 'full_end': full_end,
            'core_start': core_start, 'core_end': core_end,
            'strand': strand
        }
    return p_idx


def get_wt_base(gene_id, flank):
    clean_id = gene_id.strip()
    start_pos, end_pos, chrom, strand = None, None, None, 1
    cds_raw = []
    lines = load_gff(); genome_dict = load_genome()

    for line in lines:
        if line.startswith("#") or not line.strip(): continue
        parts = line.split("\t")
        if len(parts) < 9: continue
        if parts[2] not in ("gene", "protein_coding_gene"): continue
        attr = parts[8]
        if f"ID={clean_id};" in attr or attr.strip().endswith(f"ID={clean_id}"):
            chrom = parts[0]; start_pos = int(parts[3]); end_pos = int(parts[4])
            strand = 1 if parts[6] == "+" else -1; break

    if not chrom: return None, f"GFF 파일에서 Gene ID '{clean_id}'를 찾을 수 없습니다."

    for line in lines:
        if line.startswith("#") or not line.strip(): continue
        parts = line.split("\t")
        if len(parts) < 9: continue
        if parts[0] == chrom and parts[2] == "CDS":
            attr = parts[8]
            if f"gene_id={clean_id};" in attr or f"gene_id={clean_id}" in attr:
                cds_raw.append((int(parts[3]), int(parts[4])))

    if chrom not in genome_dict: return None, f"FASTA 파일에서 Chromosome '{chrom}'을 찾을 수 없습니다."
    full_seq = genome_dict[chrom].seq
    ext_s = max(1, start_pos - flank); ext_e = min(len(full_seq), end_pos + flank)
    sub_seq = full_seq[ext_s - 1: ext_e]
    
    if strand == -1: sub_seq = sub_seq.reverse_complement()

    cds_mapped = []
    for (s, e) in cds_raw:
        s0 = s - ext_s if strand == 1 else ext_e - e
        e0 = e - ext_s + 1 if strand == 1 else ext_e - s + 1
        cds_mapped.append((s0, e0))
        
    cds_mapped = sorted(list(set(cds_mapped)), key=lambda x: x[0])

    return {'sub_seq': sub_seq, 'cds_mapped': cds_mapped, 'gene_id': clean_id}, None


def apply_topology(record, topo):
    record.annotations["topology"] = topo; return record


def add_restriction_sites(record, seq, enz_names, gene_id=""):
    if not enz_names or str(enz_names).lower() == 'nan': return
    clean_enz = [e.strip() for e in str(enz_names).split(',') if e.strip()]
    if not clean_enz: return
    try:
        rb = RestrictionBatch(clean_enz)
        for enz, sites in rb.search(seq).items():
            for s in sites:
                record.features.append(SeqFeature(
                    FeatureLocation(s - 1, s, strand=1), type="misc_feature",
                    qualifiers={"note": [str(enz)], "label": [str(enz)]}))
    except Exception as e:
        st.warning(f"⚠️ 제한효소 인식 오류 ({gene_id}): {e}")


def add_primers_and_probes(record, seq, p_list, pb_list):
    p_idx = get_primer_coords(seq, p_list)
    not_found = [p['name'] for p in p_list if p['name'] and p['seq'] and p['name'] not in p_idx]
    if not_found:
        st.warning(f"⚠️ 서열 내에서 위치를 찾을 수 없는 프라이머: {', '.join(not_found)}")
        
    for pname, pinfo in p_idx.items():
        record.features.append(SeqFeature(
            FeatureLocation(pinfo['full_start'], pinfo['full_end'], strand=pinfo['strand']),
            type="primer_bind", qualifiers={"note": [pname], "label": [pname]}))
            
    for pb in pb_list:
        p1, p2 = pb['p1'], pb['p2']
        if p1 in p_idx and p2 in p_idx:
            coords = [p_idx[p1]['core_start'], p_idx[p1]['core_end'],
                      p_idx[p2]['core_start'], p_idx[p2]['core_end']]
            record.features.append(SeqFeature(
                FeatureLocation(min(coords), max(coords), strand=1),
                type="misc_feature", qualifiers={"note": ["Probe"], "label": ["Probe"]}))


def process_wt(gene_id, p_list, pb_list, enz_names, flank, topo):
    try:
        base, err = get_wt_base(gene_id, flank)
        if err: return None, err
        sub_seq = base['sub_seq']
        record = SeqRecord(sub_seq, id=gene_id, name=f"{gene_id}_WT",
                           annotations={"molecule_type": "DNA"})
        apply_topology(record, topo)
        
        for i, (s0, e0) in enumerate(base['cds_mapped'], 1):
            record.features.append(SeqFeature(FeatureLocation(s0, e0, strand=1), type="CDS",
                                              qualifiers={"note": [f"E{i}"], "label": [f"E{i}"]}))
                                              
        add_restriction_sites(record, sub_seq, enz_names, gene_id)
        add_primers_and_probes(record, sub_seq, p_list, pb_list)
        return record, None
    except Exception as e:
        return None, f"WT 서열 생성 오류: {e}"


def process_mutant(gene_id, p_list, pb_list, enz_names, flank,
                   ins_rec, insert_mode,
                   primer_a_name, primer_b_name, topo):
    try:
        base, err = get_wt_base(gene_id, flank)
        if err: return None, err
        sub_seq = base['sub_seq']
        ins_seq = ins_rec.seq if ins_rec else Seq("")

        def find_cut_point(primer_name, mode="end"):
            primer_entry = next((p for p in p_list if p['name'] == primer_name), None)
            if not primer_entry or not primer_entry['seq'].strip():
                return None, f"프라이머 '{primer_name}'의 서열 정보가 누락되었습니다."
            full_seq_str = primer_entry['seq'].strip().upper().replace(" ", "")
            overlap_str  = str(primer_entry.get('overlap', '')).strip().upper().replace(" ", "")
            if overlap_str and full_seq_str.startswith(overlap_str):
                search_seq = full_seq_str[len(overlap_str):]
                if not search_seq:
                    return None, f"'{primer_name}': Overlap 서열 제거 후 결합할 핵심(Core) 서열이 없습니다."
            elif overlap_str:
                st.warning(f"⚠️ {primer_name}: Overlap 서열이 프라이머의 5' 말단과 일치하지 않아 전체 서열 기준으로 검색합니다.")
                search_seq = full_seq_str
            else:
                search_seq = full_seq_str
            
            target = Seq(search_seq)
            idx = sub_seq.find(target)
            if idx != -1:
                return (idx if mode == "start" else idx + len(search_seq)), None
            idx = sub_seq.find(target.reverse_complement())
            if idx != -1:
                return (idx if mode == "start" else idx + len(search_seq)), None
            return None, f"WT 서열 내에서 '{primer_name}'의 결합 위치({search_seq[:20]}...)를 찾을 수 없습니다."

        cut_start, err_a = find_cut_point(primer_a_name, mode="end")
        if err_a: return None, err_a

        if insert_mode == "replace":
            if not primer_b_name: return None, "Replace 모드에서는 Primer B 지정이 필수입니다."
            cut_end, err_b = find_cut_point(primer_b_name, mode="start")
            if err_b: return None, err_b
            if cut_start >= cut_end:
                return None, f"Primer A의 끝({cut_start}) 위치가 Primer B의 시작({cut_end}) 위치보다 같거나 뒤에 있습니다."
            mut_seq = sub_seq[:cut_start] + ins_seq + sub_seq[cut_end:]
            ins_start_in_mut = cut_start
            ins_delta = len(ins_seq) - (cut_end - cut_start)
        else:
            cut_end = cut_start 
            mut_seq = sub_seq[:cut_start] + ins_seq + sub_seq[cut_start:]
            ins_start_in_mut = cut_start
            ins_delta = len(ins_seq)

        shifted_cds = []
        # WT 기준의 원래 인덱스(i)를 함께 저장
        for i, (s0, e0) in enumerate(base['cds_mapped'], 1):
            if e0 <= cut_start:
                shifted_cds.append((s0, e0, i))
            elif s0 >= cut_end:
                shifted_cds.append((s0 + ins_delta, e0 + ins_delta, i))
            else:
                if s0 < cut_start:
                    shifted_cds.append((s0, cut_start, i))
                if e0 > cut_end:
                    shifted_cds.append((cut_end + ins_delta, e0 + ins_delta, i))

        record = SeqRecord(mut_seq, id=gene_id, name=f"{gene_id}_MUT",
                           annotations={"molecule_type": "DNA"})
        apply_topology(record, topo)

        # 렌더링 시 오리지널 인덱스(orig_idx)를 그대로 사용하여 번호 유지
        for (s0, e0, orig_idx) in shifted_cds:
            record.features.append(SeqFeature(FeatureLocation(s0, e0, strand=1), type="CDS",
                                              qualifiers={"note": [f"E{orig_idx}"], "label": [f"E{orig_idx}"]}))

        if ins_rec:
            for feat in ins_rec.features:
                if feat.type == "source":
                    continue
                
                new_start = feat.location.start + ins_start_in_mut
                new_end = feat.location.end + ins_start_in_mut
                new_loc = FeatureLocation(new_start, new_end, strand=feat.location.strand)
                
                new_feat = SeqFeature(new_loc, type=feat.type, qualifiers=feat.qualifiers)
                record.features.append(new_feat)

        add_restriction_sites(record, mut_seq, enz_names, gene_id)
        add_primers_and_probes(record, mut_seq, p_list, pb_list)
        return record, None
    except Exception as e:
        return None, f"Mutant 서열 생성 오류: {e}"


def write_records_to_zip(zf, jobs, flank, topo, zip_structure="flat"):
    success, errors = 0, []
    total = len(jobs)
    progress = st.progress(0, text="파일 생성 처리 중...")

    for i, j in enumerate(jobs):
        progress.progress((i + 1) / total, text=f"작업 진행 중: {j['id']}")
        out_mode  = j.get('out_mode', 'wt')
        
        # 사용자가 입력한 파일명 (입력 안 했으면 기본 유전자 ID)
        custom_name = j.get('out_name', j['id'])

        prefix = f"{j['id']}/" if zip_structure == "by_gene" else ""

        # WT는 무조건 [Gene ID] WT allele 로 고정
        wt_fname  = f"{prefix}{j['id']} WT allele"
        
        # Mutant는 사용자가 지정한 이름 그대로 사용 (자동 MUT allele 꼬리표 없음)
        mut_fname = f"{prefix}{custom_name}"

        if out_mode in ('wt', 'both'):
            wt_res, wt_err = process_wt(
                j['id'], j['p_list'], j['pb_list_wt'], j['enz'], flank, topo)
            if wt_res:
                buf = StringIO(); SeqIO.write(wt_res, buf, "genbank")
                zf.writestr(f"{wt_fname}.gb", buf.getvalue())
                success += 1
            else:
                errors.append(f"❌ {j['id']} WT 작업 실패: {wt_err}")

        if out_mode in ('mut', 'both'):
            mut_res, mut_err = process_mutant(
                j['id'], j['p_list'], j['pb_list_mut'], j['enz'], flank,
                j['ins_rec'], j['mode'],
                j['pa'], j['pb'], topo)
            if mut_res:
                buf = StringIO(); SeqIO.write(mut_res, buf, "genbank")
                zf.writestr(f"{mut_fname}.gb", buf.getvalue())
                success += 1
            else:
                errors.append(f"❌ {j['id']} MUT 작업 실패: {mut_err}")

    progress.empty()
    return success, errors


# ════════════════════════════════════════════════════════════════════════════════
# ALLELE DESIGNER TAB
# ════════════════════════════════════════════════════════════════════════════════
with tab_main:
    st.header("🧬 Allele Designer")
    
    st.warning("""
    **⚠️ 사용 전 필수 안내**
    - **교차 점검 필수:** 본 도구는 연구 편의를 돕는 자동화 유틸리티입니다. 최종 생성된 시퀀스 파일 및 프라이머 결합 위치는 실험 전 반드시 **SnapGene 등의 시퀀스 뷰어를 통해 직접 교차 점검(Cross-check)** 해 주시기 바랍니다.
    - **이용 범위:** 본 도구는 **학술적 연구 목적**으로만 자유롭게 이용 가능하며, 상업적 이용은 금지됩니다.
    """)

    st.info("사이드바에서 통합 디자인 양식과 매뉴얼을 다운로드할 수 있습니다. 결과물은 WT, Mutant, 또는 둘 다 출력하도록 선택 가능합니다.")

    input_mode = st.radio("입력 방식", ["🖱️ 개별 수동 입력", "📂 엑셀 일괄 업로드"],
                          horizontal=True, key="input_mode")
    st.divider()

    if input_mode == "🖱️ 개별 수동 입력":

        col_l, col_r = st.columns(2)

        with col_l:
            st.subheader("① 기본 정보")
            gene_id_m = st.text_input("Gene ID", placeholder="예: CNAG_03701", key="m_gene").strip()
            enz_m     = st.text_input("제한효소 (쉼표 구분)", placeholder="EcoRV, BamHI", key="m_enz")

            st.subheader("② 프라이머")
            if 'm_primers' not in st.session_state: st.session_state.m_primers = []
            if st.button("프라이머 추가", key="m_add_p"):
                st.session_state.m_primers.append({"name": f"P{len(st.session_state.m_primers)+1}", "seq": ""})
            for i, p in enumerate(st.session_state.m_primers):
                with st.expander(f"Primer: {p['name']}", expanded=True):
                    p['name']    = st.text_input("이름",   value=p['name'], key=f"m_pn_{i}").strip()
                    p['seq']     = st.text_input("서열",   value=p['seq'],  key=f"m_ps_{i}")
                    p['overlap'] = st.text_input(
                        "Overlap (5' 마커 결합 서열, 없으면 공란)",
                        value=p.get('overlap', ''), key=f"m_pov_{i}",
                        placeholder="예: CACTCGAATCCTGCATGC"
                    ).strip()
            if st.button("프라이머 초기화", key="m_rst_p"):
                st.session_state.m_primers = []; st.rerun()

            st.subheader("③-1. Southern Probe (WT)")
            if 'm_probes_wt' not in st.session_state: st.session_state.m_probes_wt = []
            if st.button("WT Probe 추가", key="m_add_pb_wt"):
                st.session_state.m_probes_wt.append({"p1": "", "p2": ""})
            for i, pb in enumerate(st.session_state.m_probes_wt):
                c1, c2 = st.columns(2)
                with c1: pb['p1'] = st.text_input("시작 프라이머", value=pb['p1'], key=f"m_pb1_wt_{i}").strip()
                with c2: pb['p2'] = st.text_input("끝 프라이머",   value=pb['p2'], key=f"m_pb2_wt_{i}").strip()
            if st.button("WT Probe 초기화", key="m_rst_pb_wt"):
                st.session_state.m_probes_wt = []; st.rerun()

            st.subheader("③-2. Southern Probe (MUT)")
            if 'm_probes_mut' not in st.session_state: st.session_state.m_probes_mut = []
            if st.button("MUT Probe 추가", key="m_add_pb_mut"):
                st.session_state.m_probes_mut.append({"p1": "", "p2": ""})
            for i, pb in enumerate(st.session_state.m_probes_mut):
                c1, c2 = st.columns(2)
                with c1: pb['p1'] = st.text_input("시작 프라이머", value=pb['p1'], key=f"m_pb1_mut_{i}").strip()
                with c2: pb['p2'] = st.text_input("끝 프라이머",   value=pb['p2'], key=f"m_pb2_mut_{i}").strip()
            if st.button("MUT Probe 초기화", key="m_rst_pb_mut"):
                st.session_state.m_probes_mut = []; st.rerun()

        with col_r:
            st.subheader("④ 출력 모드")
            out_mode_m = st.radio(
                "출력 대상",
                options=["wt", "mut", "both"],
                format_func=lambda x: {"wt": "WT 서열만", "mut": "Mutant 서열만", "both": "WT + Mutant 모두"}[x],
                key="m_out_mode", horizontal=True
            )
            out_name_m = st.text_input(
                "Mutant 출력 파일명 (WT는 'Gene ID WT allele'로 고정됨)",
                key="m_out_name", placeholder="예: CNAG_03701_NAT"
            ).strip()

            if out_mode_m in ("mut", "both"):
                st.subheader("⑤ 삽입 방식 설정")
                ins_mode_m = st.radio(
                    "삽입 형태",
                    options=["insert", "replace"],
                    format_func=lambda x: (
                        "➕ Insert — Primer A 3' 말단 직후 삽입" if x == "insert"
                        else "🔄 Replace — Primer A 3' 말단부터 Primer B 5' 시작 구간 치환"),
                    key="m_ins_mode"
                )
                pnames = [p['name'] for p in st.session_state.m_primers if p['name']]
                c1, c2 = st.columns(2)
                with c1:
                    pa_m    = st.selectbox("Primer A", pnames or ["—"], key="m_pa")
                    pa_ov_m = st.text_input("Primer A Overlap (5' 마커 결합 서열)",
                                            key="m_pa_ov", placeholder="ATGCATGC...").strip()
                with c2:
                    if ins_mode_m == "replace":
                        pb_m    = st.selectbox("Primer B", pnames or ["—"], key="m_pb")
                        pb_ov_m = st.text_input("Primer B Overlap (5' 마커 결합 서열)",
                                                key="m_pb_ov", placeholder="GCTAGCTA...").strip()
                    else:
                        pb_m = None; pb_ov_m = ""

                st.subheader("⑥ 삽입 서열 (GenBank)")
                st.info("업로드한 마커 파일 내의 주석(Feature) 정보는 자동으로 보존되어 이식됩니다.")
                ins_gb = st.file_uploader("삽입할 .gb / .gbk 파일 업로드", type=["gb","gbk"], key="m_ins_gb")
                ins_rec_m = None
                if ins_gb:
                    try:
                        ins_rec_m = SeqIO.read(StringIO(ins_gb.read().decode("utf-8")), "genbank")
                        _dn = ins_rec_m.name if ins_rec_m.name and ins_rec_m.name != '.' else ins_rec_m.id
                        st.success(f"✅ {ins_rec_m.id} ({len(ins_rec_m.seq)} bp) 정상 로드")
                    except Exception as e:
                        st.error(f"GB 파일 판독 오류: {e}")

        if st.button("GenBank 파일 생성 🚀", key="m_run"):
            if not gene_id_m:
                st.warning("Gene ID를 입력해 주세요.")
            elif out_mode_m in ("mut", "both") and not ins_rec_m:
                st.warning("삽입할 마커(GenBank) 파일을 업로드해 주세요.")
            else:
                _pa     = st.session_state.get('m_pa', '')
                _pb     = st.session_state.get('m_pb', None)
                _p_list = [dict(p) for p in st.session_state.m_primers]
                job = {
                    'id': gene_id_m,
                    'p_list': _p_list,
                    'pb_list_wt': st.session_state.m_probes_wt,
                    'pb_list_mut': st.session_state.m_probes_mut,
                    'enz': enz_m,
                    'out_mode': out_mode_m,
                    'out_name': out_name_m or gene_id_m,
                    'mode': st.session_state.get('m_ins_mode', 'insert'),
                    'pa': _pa,
                    'pb': _pb,
                    'ins_rec': ins_rec_m, 
                }
                zip_buf = BytesIO()
                with zipfile.ZipFile(zip_buf, "a") as zf:
                    success, errors = write_records_to_zip(zf, [job], flank_size, topology, zip_structure)
                if success:
                    now = datetime.now().strftime("%Y%m%d_%H%M")
                    st.success(f"✅ 총 {success}개의 파일 생성이 완료되었습니다. (Topology: {topology})")
                    st.download_button(f"📥 다운로드 ({now})", zip_buf.getvalue(), f"Results_{now}.zip")
                for e in errors: st.write(e)

    else:
        st.markdown("""
**통합 양식 시트 구성 안내** (사이드바에서 다운로드 가능):

| 시트명 | 주요 기입 항목 |
|------|-----------|
| `Primers` | Gene ID / Primer Name / Sequence |
| `Probes_WT` | Gene ID / Probe Start Primer / Probe End Primer |
| `Probes_MUT` | Gene ID / Probe Start Primer / Probe End Primer |
| `Enzymes` | Gene ID / Enzymes |
| `Jobs` | Gene ID / **Output Mode** / Insert Mode / Primer A / Primer A Overlap / Primer B / Primer B Overlap / Insert GB Filename / Output Filename |

* `Output Mode`: `wt` / `mut` / `both` (WT 서열만 출력 시 삽입 관련 옵션은 무시됩니다.)
        """)

        excel_file = st.file_uploader("작성된 통합 엑셀 파일 업로드", type=["xlsx"], key="batch_excel")
        gb_uploads = st.file_uploader("삽입 서열 파일 업로드 (다중 선택 가능)",
                                      type=["gb","gbk"], accept_multiple_files=True, key="batch_gb")
        gb_dict = {}
        if gb_uploads:
            for f in gb_uploads:
                try:
                    rec = SeqIO.read(StringIO(f.read().decode("utf-8")), "genbank")
                    gb_dict[f.name] = rec
                    st.success(f"✅ {f.name}: {rec.id} ({len(rec.seq)} bp) 정상 로드")
                except Exception as e:
                    st.error(f"❌ {f.name} 파일 로드 실패: {e}")

        if st.button("GenBank 파일 생성 🚀", key="batch_run"):
            if not excel_file:
                st.warning("엑셀 파일을 업로드해 주세요.")
            else:
                try:
                    df_jobs    = pd.read_excel(excel_file, sheet_name='Jobs')
                    df_primers = pd.read_excel(excel_file, sheet_name='Primers')
                    df_enz     = pd.read_excel(excel_file, sheet_name='Enzymes')
                    
                    try:
                        df_probes_wt = pd.read_excel(excel_file, sheet_name='Probes_WT')
                    except Exception:
                        df_probes_wt = pd.DataFrame(columns=['Gene ID', 'Probe Start Primer', 'Probe End Primer'])
                    
                    try:
                        df_probes_mut = pd.read_excel(excel_file, sheet_name='Probes_MUT')
                    except Exception:
                        df_probes_mut = pd.DataFrame(columns=['Gene ID', 'Probe Start Primer', 'Probe End Primer'])

                    def clean(v): return '' if str(v).strip().lower() == 'nan' else str(v).strip()

                    jobs = []
                    for _, row in df_jobs.iterrows():
                        gene     = clean(row['Gene ID'])
                        out_mode = clean(row.get('Output Mode', 'wt')).lower()
                        out_mode = out_mode if out_mode in ('wt','mut','both') else 'wt'
                        ins_mode = clean(row.get('Insert Mode', 'insert')).lower()
                        pa       = clean(row.get('Primer A', ''))
                        pa_ov    = clean(row.get('Primer A Overlap', ''))
                        pb       = clean(row.get('Primer B', ''))
                        pb_ov    = clean(row.get('Primer B Overlap', ''))
                        gb_fname = clean(row.get('Insert GB Filename', ''))
                        out_nm   = clean(row.get('Output Filename', '')) or gene

                        p_list = [
                            {
                                'name': clean(r['Primer Name']),
                                'seq':  clean(r['Sequence']),
                                'overlap': pa_ov if clean(r['Primer Name']) == pa else (pb_ov if clean(r['Primer Name']) == pb else "")
                            }
                            for _, r in df_primers[df_primers['Gene ID'] == gene].iterrows()
                        ]
                        
                        pb_list_wt = [
                            {'p1': clean(r['Probe Start Primer']), 'p2': clean(r['Probe End Primer'])}
                            for _, r in df_probes_wt[df_probes_wt['Gene ID'] == gene].iterrows()
                        ]
                        
                        pb_list_mut = [
                            {'p1': clean(r['Probe Start Primer']), 'p2': clean(r['Probe End Primer'])}
                            for _, r in df_probes_mut[df_probes_mut['Gene ID'] == gene].iterrows()
                        ]
                        
                        enz = df_enz[df_enz['Gene ID'] == gene]['Enzymes'].iloc[0] \
                              if gene in df_enz['Gene ID'].values else ""

                        ins_rec  = None
                        if out_mode in ('mut', 'both'):
                            if not gb_fname or gb_fname not in gb_dict:
                                st.error(f"❌ {gene}: 대상 마커 파일('{gb_fname}')이 업로드되지 않았습니다.")
                                continue
                            ins_rec = gb_dict[gb_fname]

                        jobs.append({
                            'id': gene, 'p_list': p_list, 'pb_list_wt': pb_list_wt, 'pb_list_mut': pb_list_mut, 'enz': enz,
                            'out_mode': out_mode, 'out_name': out_nm,
                            'mode': ins_mode, 'pa': pa,
                            'pb': pb if ins_mode == 'replace' else None,
                            'ins_rec': ins_rec,
                        })

                    if jobs:
                        zip_buf = BytesIO()
                        with zipfile.ZipFile(zip_buf, "a") as zf:
                            success, errors = write_records_to_zip(zf, jobs, flank_size, topology, zip_structure)
                        if success:
                            now = datetime.now().strftime("%Y%m%d_%H%M")
                            st.success(f"✅ 총 {success}개의 파일 생성이 완료되었습니다. (Topology: {topology})")
                            st.download_button(f"📥 다운로드 ({now})", zip_buf.getvalue(), f"Results_{now}.zip")
                        for e in errors: st.write(e)

                except Exception as ex:
                    st.error(f"엑셀 데이터 처리 오류: {ex}")


# ════════════════════════════════════════════════════════════════════════════════
# GB → DNA 변환 탭
# ════════════════════════════════════════════════════════════════════════════════
with tab_conv:
    st.header("🔄 GB → DNA 확장자 일괄 변환")
    st.info("업로드된 `.gb` 또는 `.gbk` 파일의 확장자를 `.dna`로 일괄 변환하여 ZIP 압축 파일로 제공합니다.")
    conv_files = st.file_uploader("변환 대상 파일 선택", type=["gb","gbk"],
                                  accept_multiple_files=True, key="conv_files")
    if conv_files:
        st.write(f"업로드된 파일 총 {len(conv_files)}개:")
        for f in conv_files:
            st.write(f"  • {f.name} → {f.name.rsplit('.', 1)[0]}.dna")
        if st.button("🔄 파일 변환 및 ZIP 다운로드", key="conv_run"):
            zip_buf = BytesIO()
            with zipfile.ZipFile(zip_buf, "w") as zf:
                for f in conv_files:
                    zf.writestr(f.name.rsplit('.', 1)[0] + ".dna", f.read())
            now = datetime.now().strftime("%Y%m%d_%H%M")
            st.success(f"✅ {len(conv_files)}개 파일의 포맷 변환이 완료되었습니다.")
            st.download_button(f"📥 변환된 파일 다운로드 ({now})", zip_buf.getvalue(), f"DNA_files_{now}.zip")